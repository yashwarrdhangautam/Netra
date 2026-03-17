"""
netra/reports/word_report.py
Professional .docx penetration test report generator.
Produces a fully formatted Word document with:
  - Title page, executive summary, methodology
  - Severity-coloured findings table
  - Per-finding detail sections with screenshots
  - Attack chain appendix
  - Remediation roadmap

Styling:
  Font: Calibri 11pt body, 14pt H1, 12pt H2
  Heading colour: #1F3864 (navy), Accent: #2E75B6 (blue)
  Severity: Critical=#FF0000, High=#FF6600, Med=#FFC000, Low=#92D050
"""

import re
from pathlib import Path
from datetime import datetime
from typing import Optional

# Colour constants
COL_NAVY     = "1F3864"
COL_BLUE     = "2E75B6"
COL_CRITICAL = "FF0000"
COL_HIGH     = "FF6600"
COL_MEDIUM   = "FFC000"
COL_LOW      = "92D050"
COL_INFO     = "888888"
COL_WHITE    = "FFFFFF"
COL_LGREY    = "F2F2F2"

SEV_COLORS = {
    "critical": COL_CRITICAL,
    "high":     COL_HIGH,
    "medium":   COL_MEDIUM,
    "low":      COL_LOW,
    "info":     COL_INFO,
}


def generate_word_report(ctx: dict, reports_dir: Path) -> Path:
    """
    Generate a fully formatted .docx penetration test report.

    Args:
        ctx:         Report context dict with keys: scan_id, findings, assets,
                     chains, stats, risk_score, risk_grade, client, engagement,
                     date, operator.
        reports_dir: Directory to save the output file.

    Returns:
        Path to the generated .docx file.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise ImportError("python-docx required. Run: pip3 install python-docx --break-system-packages")

    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── Document styles ───────────────────────────────────────────────
    _apply_styles(doc)

    # ── Title page ────────────────────────────────────────────────────
    _add_title_page(doc, ctx)

    # ── Table of contents placeholder ─────────────────────────────────
    doc.add_page_break()
    _heading(doc, "Table of Contents", 1)
    p = doc.add_paragraph("[Update field: right-click → Update Field]")
    p.style.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ── Executive summary ─────────────────────────────────────────────
    doc.add_page_break()
    _heading(doc, "Executive Summary", 1)
    _add_executive_summary(doc, ctx)

    # ── Scope & Methodology ───────────────────────────────────────────
    _heading(doc, "Scope & Methodology", 1)
    _add_methodology(doc, ctx)

    # ── Risk scorecard ────────────────────────────────────────────────
    _heading(doc, "Risk Scorecard", 1)
    _add_risk_scorecard(doc, ctx)

    # ── Findings overview table ───────────────────────────────────────
    _heading(doc, "Findings Overview", 1)
    _add_findings_table(doc, ctx)

    # ── Detailed findings ─────────────────────────────────────────────
    _heading(doc, "Detailed Findings", 1)
    for finding in ctx.get("findings", []):
        _add_finding_detail(doc, finding)

    # ── Attack chains ─────────────────────────────────────────────────
    chains = ctx.get("chains", [])
    if chains:
        _heading(doc, "Attack Chain Analysis", 1)
        _add_attack_chains(doc, chains, ctx.get("findings", []))

    # ── Remediation roadmap ───────────────────────────────────────────
    _heading(doc, "Remediation Roadmap", 1)
    _add_remediation_roadmap(doc, ctx)

    # ── Appendix: Asset inventory ─────────────────────────────────────
    assets = ctx.get("assets", [])
    if assets:
        _heading(doc, "Appendix A: Asset Inventory", 1)
        _add_asset_table(doc, assets)

    # ── Save ──────────────────────────────────────────────────────────
    filename = _report_filename(ctx, "docx")
    out_path = reports_dir / filename
    doc.save(str(out_path))
    return out_path


# ── Private helpers ───────────────────────────────────────────────────

def _apply_styles(doc) -> None:
    """Configure default document font and paragraph styles."""
    from docx.shared import Pt
    style            = doc.styles["Normal"]
    style.font.name  = "Calibri"
    style.font.size  = Pt(11)


def _heading(doc, text: str, level: int) -> None:
    """Add a styled heading paragraph."""
    from docx.shared import Pt, RGBColor
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    run.font.name = "Calibri"
    run.font.size = Pt(14 if level == 1 else 12)


def _add_title_page(doc, ctx: dict) -> None:
    """Add the report title page."""
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Report title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("PENETRATION TEST REPORT")
    run.font.name  = "Calibri"
    run.font.size  = Pt(28)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    doc.add_paragraph()

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run("NETRA नेत्र — The Third Eye of Security")
    run.font.name  = "Calibri"
    run.font.size  = Pt(14)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    doc.add_paragraph()
    doc.add_paragraph()

    # Meta table
    meta = doc.add_table(rows=6, cols=2)
    meta.style = "Table Grid"
    fields = [
        ("Client",       ctx.get("client", "Confidential")),
        ("Engagement",   ctx.get("engagement", "Security Assessment")),
        ("Date",         ctx.get("date", datetime.now().strftime("%Y-%m-%d"))),
        ("Operator",     ctx.get("operator", "Security Team")),
        ("Risk Grade",   ctx.get("risk_grade", "?") + f"  ({ctx.get('risk_score', 0)}/100)"),
        ("Scan ID",      ctx.get("scan_id", "N/A")),
    ]
    for i, (label, value) in enumerate(fields):
        row = meta.rows[i]
        row.cells[0].text = label
        row.cells[1].text = str(value)
        row.cells[0].paragraphs[0].runs[0].font.bold = True

    doc.add_page_break()


def _add_executive_summary(doc, ctx: dict) -> None:
    """Add executive summary section."""
    from docx.shared import Pt

    # Try AI-generated summary
    ai_summary = ctx.get("ai_summary", "")
    if not ai_summary:
        try:
            from netra.ai_brain.narrative import generate_executive_summary
            ai_summary = generate_executive_summary(ctx)
        except Exception:
            ai_summary = ""

    if ai_summary:
        for para in ai_summary.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
    else:
        # Fallback
        stats  = ctx.get("stats", {})
        total  = sum(stats.values())
        grade  = ctx.get("risk_grade", "?")
        score  = ctx.get("risk_score", 0)
        doc.add_paragraph(
            f"This security assessment identified {total} vulnerabilities across the assessed scope, "
            f"resulting in an overall risk grade of {grade} ({score}/100). "
            f"The findings represent a range of risks that require prioritised remediation."
        )

    doc.add_paragraph()


def _add_methodology(doc, ctx: dict) -> None:
    """Add scope and methodology section."""
    assets = ctx.get("assets", [])
    doc.add_paragraph(
        "The assessment was conducted using the NETRA security platform, "
        "employing a combination of automated scanning and manual verification. "
        "The following methodology phases were executed:"
    )

    phases = [
        ("Reconnaissance",     "OSINT, subdomain enumeration, technology fingerprinting"),
        ("Asset Discovery",    "HTTP probing, port scanning, web crawling"),
        ("Vulnerability Scan", "Nuclei template scanning, CVE checks"),
        ("Exploitation",       "Manual verification, injection testing, auth bypass"),
        ("AI Analysis",        "Multi-persona consensus validation, false positive filtering"),
        ("Reporting",          "CVSS scoring, MITRE mapping, chain discovery"),
    ]

    table = doc.add_table(rows=len(phases) + 1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0]
    hdr.cells[0].text = "Phase"
    hdr.cells[1].text = "Activities"
    for cell in hdr.cells:
        cell.paragraphs[0].runs[0].font.bold = True

    for i, (phase, activities) in enumerate(phases, 1):
        table.rows[i].cells[0].text = phase
        table.rows[i].cells[1].text = activities

    doc.add_paragraph()
    doc.add_paragraph(f"Total assets discovered: {len(assets)}")


def _add_risk_scorecard(doc, ctx: dict) -> None:
    """Add risk scorecard with severity counts."""
    from docx.shared import RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    stats = ctx.get("stats", {})
    score = ctx.get("risk_score", 0)
    grade = ctx.get("risk_grade", "?")

    p = doc.add_paragraph()
    run = p.add_run(f"Risk Score: {score}/100  |  Grade: {grade}")
    run.font.bold = True
    run.font.size = __import__("docx").shared.Pt(14)

    doc.add_paragraph()

    cols = ["Severity", "Count", "Risk Level"]
    rows_data = [
        ("Critical", stats.get("critical", 0), "Immediate action required"),
        ("High",     stats.get("high", 0),     "Address within 24-48 hours"),
        ("Medium",   stats.get("medium", 0),   "Address within current sprint"),
        ("Low",      stats.get("low", 0),      "Address in backlog"),
        ("Info",     stats.get("info", 0),     "Informational — no action required"),
    ]

    table = doc.add_table(rows=len(rows_data) + 1, cols=3)
    table.style = "Table Grid"

    for j, col in enumerate(cols):
        cell = table.rows[0].cells[j]
        cell.text = col
        cell.paragraphs[0].runs[0].font.bold = True

    sev_cols = {
        "critical": COL_CRITICAL, "high": COL_HIGH,
        "medium": COL_MEDIUM, "low": COL_LOW, "info": COL_INFO,
    }

    for i, (sev, count, risk) in enumerate(rows_data, 1):
        row = table.rows[i]
        row.cells[0].text = sev
        row.cells[1].text = str(count)
        row.cells[2].text = risk
        # Colour severity cell
        hex_col = sev_cols.get(sev.lower(), COL_INFO)
        _set_cell_bg(row.cells[0], hex_col)

    doc.add_paragraph()


def _add_findings_table(doc, ctx: dict) -> None:
    """Add a summary table of all findings."""
    from docx.shared import Pt

    findings = ctx.get("findings", [])
    if not findings:
        doc.add_paragraph("No findings identified.")
        return

    headers = ["#", "Title", "Severity", "CVSS", "Host", "Status"]
    table   = doc.add_table(rows=len(findings) + 1, cols=len(headers))
    table.style = "Table Grid"

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        _set_cell_bg(cell, COL_NAVY)
        cell.paragraphs[0].runs[0].font.color.rgb = __import__("docx").shared.RGBColor(
            0xFF, 0xFF, 0xFF
        )

    for i, f in enumerate(findings, 1):
        row = table.rows[i]
        row.cells[0].text = str(i)
        row.cells[1].text = (f.get("title") or "")[:60]
        row.cells[2].text = (f.get("severity") or "").capitalize()
        row.cells[3].text = str(f.get("cvss_score") or "N/A")
        row.cells[4].text = (f.get("host") or "")[:40]
        row.cells[5].text = (f.get("status") or "open").capitalize()

        sev_col = SEV_COLORS.get((f.get("severity") or "info").lower(), COL_INFO)
        _set_cell_bg(row.cells[2], sev_col)

    doc.add_paragraph()


def _add_finding_detail(doc, f: dict) -> None:
    """Add a detailed section for one finding."""
    from docx.shared import Pt, RGBColor

    sev = (f.get("severity") or "info").lower()
    _heading(doc, f.get("title", "Unnamed Finding"), 2)

    # Meta table
    meta_table = doc.add_table(rows=8, cols=2)
    meta_table.style = "Table Grid"

    meta_rows = [
        ("Severity",   (f.get("severity") or "").capitalize()),
        ("CVSS Score", str(f.get("cvss_score") or "N/A")),
        ("CVE",        f.get("cve_id") or "N/A"),
        ("CWE",        f.get("cwe_id") or "N/A"),
        ("Host",       f.get("host") or "N/A"),
        ("URL",        (f.get("url") or "")[:80]),
        ("MITRE",      f.get("mitre_technique") or "N/A"),
        ("Status",     (f.get("status") or "open").capitalize()),
    ]

    for i, (label, value) in enumerate(meta_rows):
        row = meta_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        if label == "Severity" and value:
            _set_cell_bg(row.cells[1], SEV_COLORS.get(value.lower(), COL_INFO))

    doc.add_paragraph()

    # Description
    if f.get("description"):
        _heading(doc, "Description", 3)
        doc.add_paragraph(f["description"][:800])

    # Impact
    if f.get("impact"):
        _heading(doc, "Impact", 3)
        doc.add_paragraph(f["impact"][:400])

    # AI Narrative
    if f.get("ai_narrative"):
        _heading(doc, "Analysis", 3)
        doc.add_paragraph(f["ai_narrative"][:600])

    # Evidence
    if f.get("evidence") or f.get("poc_command"):
        _heading(doc, "Evidence / PoC", 3)
        if f.get("poc_command"):
            p = doc.add_paragraph()
            run = p.add_run(f.get("poc_command", "")[:200])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        if f.get("evidence"):
            p = doc.add_paragraph()
            run = p.add_run(f.get("evidence", "")[:400])
            run.font.name = "Courier New"
            run.font.size = Pt(9)

    # Remediation
    if f.get("remediation"):
        _heading(doc, "Remediation", 3)
        doc.add_paragraph(f["remediation"][:400])

    doc.add_paragraph()


def _add_attack_chains(doc, chains: list, all_findings: list) -> None:
    """Add attack chain analysis section."""
    id_map = {f["id"]: f for f in all_findings}

    for i, chain in enumerate(chains[:10], 1):
        _heading(doc, f"Chain {i}: CVSS {chain.get('combined_cvss', 'N/A')}", 2)

        p = doc.add_paragraph()
        p.add_run("MITRE Sequence: ").bold = True
        p.add_run(chain.get("mitre_sequence", "N/A"))

        if chain.get("narrative"):
            doc.add_paragraph(chain["narrative"][:500])

        # Chain steps
        import json as _json
        try:
            node_ids = _json.loads(chain.get("nodes", "[]")) if isinstance(chain.get("nodes"), str) else chain.get("nodes", [])
        except Exception:
            node_ids = []

        for j, fid in enumerate(node_ids, 1):
            f = id_map.get(fid)
            if f:
                doc.add_paragraph(
                    f"  Step {j}: {f.get('title')} [{f.get('severity', '').upper()}] "
                    f"on {f.get('host')} (CVSS: {f.get('cvss_score', 'N/A')})"
                )

        doc.add_paragraph()


def _add_remediation_roadmap(doc, ctx: dict) -> None:
    """Add a prioritised remediation roadmap table."""
    try:
        from netra.ai_brain.analyzer import build_remediation_roadmap
        roadmap = build_remediation_roadmap(ctx.get("findings", []))
    except Exception:
        roadmap = []

    if not roadmap:
        doc.add_paragraph("No findings to remediate.")
        return

    headers = ["Priority", "Severity", "Count", "Effort", "Sample Findings"]
    table   = doc.add_table(rows=len(roadmap) + 1, cols=len(headers))
    table.style = "Table Grid"

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True

    for i, item in enumerate(roadmap, 1):
        row = table.rows[i]
        row.cells[0].text = str(item["priority"])
        row.cells[1].text = item["severity"].capitalize()
        row.cells[2].text = str(item["count"])
        row.cells[3].text = item["effort"]
        row.cells[4].text = ", ".join(item["top_titles"][:2])[:60]
        _set_cell_bg(row.cells[1], SEV_COLORS.get(item["severity"], COL_INFO))

    doc.add_paragraph()


def _add_asset_table(doc, assets: list) -> None:
    """Add asset inventory appendix table."""
    if not assets:
        doc.add_paragraph("No assets discovered.")
        return

    headers = ["Asset", "Type", "Live", "Tech Stack", "Ports"]
    table   = doc.add_table(rows=min(len(assets) + 1, 101), cols=len(headers))
    table.style = "Table Grid"

    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
        table.rows[0].cells[j].paragraphs[0].runs[0].font.bold = True

    import json as _json
    for i, a in enumerate(assets[:100], 1):
        row = table.rows[i]
        row.cells[0].text = (a.get("value") or "")[:50]
        row.cells[1].text = a.get("asset_type") or ""
        row.cells[2].text = "✓" if a.get("is_live") else ""
        try:
            ts = _json.loads(a.get("tech_stack") or "{}")
            row.cells[3].text = ", ".join(f"{k}:{v}" for k, v in ts.items())[:50]
        except Exception:
            row.cells[3].text = ""
        try:
            ports = _json.loads(a.get("ports") or "[]")
            row.cells[4].text = ", ".join(str(p) for p in ports[:8])
        except Exception:
            row.cells[4].text = ""

    doc.add_paragraph()


def _set_cell_bg(cell, hex_color: str) -> None:
    """Set background colour of a table cell."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tc_pr = cell._tc.get_or_add_tcPr()
    shd   = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _report_filename(ctx: dict, ext: str) -> str:
    """Generate a standardised report filename."""
    scan_id = ctx.get("scan_id", "unknown")
    date    = ctx.get("date", datetime.now().strftime("%Y%m%d")).replace("-", "")
    module  = "vapt"
    target  = re.sub(r"[^a-zA-Z0-9]", "_", ctx.get("engagement", scan_id))[:20]
    return f"NETRA_{module}_{target}_{date}.{ext}"
