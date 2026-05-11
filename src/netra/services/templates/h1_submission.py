"""HackerOne-style Word submission packet template."""
from __future__ import annotations

import json
from pathlib import Path
from typing import Any


def build_h1_docx(finding: Any, submission: Any, output_path: Path) -> Path:
    """Build a Word document for a HackerOne-style submission."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import RGBColor

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    title = doc.add_heading(submission.title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sev = str(submission.severity).upper()
    p = doc.add_paragraph()
    run = p.add_run(f"Severity: {sev}")
    run.bold = True
    if sev == "CRITICAL":
        run.font.color.rgb = RGBColor(180, 0, 0)
    elif sev == "HIGH":
        run.font.color.rgb = RGBColor(220, 90, 0)

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(finding.description or submission.title)

    doc.add_heading("Steps to Reproduce", level=1)
    evidence = finding.evidence or {}
    steps = evidence.get("steps_to_reproduce") if isinstance(evidence, dict) else None
    if not steps:
        steps = [
            f"Navigate to {finding.url or 'the affected asset'}.",
            "Reproduce the vulnerable behavior using the evidence below.",
            "Observe the security impact described in this packet.",
        ]
    for step in steps:
        doc.add_paragraph(str(step), style="List Number")

    doc.add_heading("Impact", level=1)
    bounty = (finding.ai_analysis or {}).get("bounty_hunter", {})
    impact = bounty.get("rationale") or (finding.ai_analysis or {}).get("attacker", {}).get("business_impact")
    doc.add_paragraph(impact or "Impact requires operator confirmation before submission.")

    doc.add_heading("Suggested Fix", level=1)
    defender = (finding.ai_analysis or {}).get("defender", {})
    doc.add_paragraph(defender.get("immediate_fix") or finding.remediation or "Apply the vendor-recommended remediation and verify with a regression test.")

    doc.add_heading("References", level=1)
    refs = evidence.get("references", []) if isinstance(evidence, dict) else []
    if refs:
        for ref in refs:
            doc.add_paragraph(str(ref), style="List Bullet")
    else:
        doc.add_paragraph("No external references attached.")

    doc.add_heading("CVSS", level=1)
    doc.add_paragraph(f"Vector: {submission.cvss_vector or finding.cvss_vector or 'N/A'}")
    if finding.cvss_score is not None:
        doc.add_paragraph(f"Score: {finding.cvss_score}")

    doc.add_heading("Evidence", level=1)
    doc.add_paragraph(json.dumps(evidence, indent=2, default=str)[:5000])

    doc.save(str(output_path))
    return output_path

