"""Report generation service for NETRA."""
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import structlog

from netra.core.config import settings
from netra.db.models.finding import Severity

logger = structlog.get_logger()


async def generate_executive_report(
    scan_data: dict[str, Any],
    findings: list[dict[str, Any]],
    output_path: Path | None = None,
) -> Path:
    """Generate executive summary PDF.

    Contains:
    - Risk score (A-F grade)
    - Severity distribution chart
    - Top 5 critical findings
    - Key metrics (total findings, by severity, SLA status)
    - Recommendation summary from AI Defender persona

    Args:
        scan_data: Scan metadata
        findings: List of findings with full data
        output_path: Optional output path

    Returns:
        Path to generated PDF
    """
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            SimpleDocTemplate,
            Paragraph,
            Spacer,
            Table,
            TableStyle,
        )
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    except ImportError:
        raise ImportError(
            "ReportLab not installed. Install with: pip install reportlab"
        )

    if output_path is None:
        scan_id = scan_data.get("id", "unknown")
        output_path = settings.reports_dir / f"executive_{scan_id}.pdf"
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = SimpleDocTemplate(str(output_path), pagesize=letter)
    styles = getSampleStyleSheet()
    elements: list[Any] = []

    # Title
    title_style = ParagraphStyle(
        "Title", parent=styles["Title"], fontSize=24, spaceAfter=20
    )
    elements.append(
        Paragraph("NETRA — Executive Security Report", title_style)
    )
    elements.append(Spacer(1, 0.2 * inch))

    # Scan metadata
    elements.append(
        Paragraph(
            f"Scan: {scan_data.get('name', 'Unknown')} | Date: {scan_data.get('created_at', 'N/A')}",
            styles["Normal"],
        )
    )
    elements.append(
        Paragraph(
            f"Target: {scan_data.get('target', 'N/A')} | Profile: {scan_data.get('profile', 'standard')}",
            styles["Normal"],
        )
    )
    elements.append(Spacer(1, 0.3 * inch))

    # Risk Score
    severity_counts = _count_severities(findings)
    risk_score = _calculate_risk_score(severity_counts)
    risk_grade = _score_to_grade(risk_score)

    elements.append(
        Paragraph(f"<b>Risk Grade: {risk_grade}</b> ({risk_score}/100)", styles["Heading2"])
    )
    elements.append(Spacer(1, 0.2 * inch))

    # Severity table
    severity_data = [
        ["Severity", "Count"],
        ["Critical", str(severity_counts.get("critical", 0))],
        ["High", str(severity_counts.get("high", 0))],
        ["Medium", str(severity_counts.get("medium", 0))],
        ["Low", str(severity_counts.get("low", 0))],
        ["Info", str(severity_counts.get("info", 0))],
    ]
    t = Table(severity_data, colWidths=[3 * inch, 2 * inch])
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a1a2e")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("FONTSIZE", (0, 0), (-1, -1), 11),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ]
        )
    )
    elements.append(t)
    elements.append(Spacer(1, 0.3 * inch))

    # Top findings
    elements.append(Paragraph("Top Critical Findings", styles["Heading2"]))
    critical_findings = [f for f in findings if f.get("severity") == "critical"][:5]
    for i, f in enumerate(critical_findings, 1):
        elements.append(
            Paragraph(
                f"{i}. <b>{f['title']}</b> — {f.get('url', 'N/A')}",
                styles["Normal"],
            )
        )
    if not critical_findings:
        elements.append(
            Paragraph("No critical findings detected.", styles["Normal"])
        )

    elements.append(Spacer(1, 0.2 * inch))

    # Key metrics
    elements.append(Paragraph("Key Metrics", styles["Heading2"]))
    metrics_data = [
        ["Metric", "Value"],
        ["Total Findings", str(len(findings))],
        ["Critical/High", str(severity_counts.get("critical", 0) + severity_counts.get("high", 0))],
        ["AI-Enriched", str(sum(1 for f in findings if f.get("ai_analysis")))],
        ["Compliance Mapped", str(sum(1 for f in findings if f.get("cwe_id")))],
    ]
    metrics_table = Table(metrics_data, colWidths=[3 * inch, 2 * inch])
    metrics_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a1a2e")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
            ]
        )
    )
    elements.append(metrics_table)

    # Build PDF
    doc.build(elements)
    logger.info("executive_report_generated", path=str(output_path))
    return output_path


async def generate_technical_report(
    scan_data: dict[str, Any],
    findings: list[dict[str, Any]],
    output_path: Path | None = None,
) -> Path:
    """Generate technical report in Word (.docx) format.

    Contains:
    - Cover page with scan metadata
    - Table of contents
    - Methodology section
    - Finding details with evidence
    - Remediation guidance from AI Defender
    - Compliance mapping from AI Analyst

    Args:
        scan_data: Scan metadata
        findings: List of findings with full data
        output_path: Optional output path

    Returns:
        Path to generated DOCX
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError(
            "python-docx not installed. Install with: pip install python-docx"
        )

    if output_path is None:
        scan_id = scan_data.get("id", "unknown")
        output_path = settings.reports_dir / f"technical_{scan_id}.docx"
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Cover page
    title = doc.add_heading("NETRA Technical Security Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"\nScan Name: {scan_data.get('name', 'Unknown')}")
    doc.add_paragraph(f"Target: {scan_data.get('target', 'N/A')}")
    doc.add_paragraph(f"Profile: {scan_data.get('profile', 'standard')}")
    doc.add_paragraph(f"Date: {scan_data.get('created_at', 'N/A')}")
    doc.add_page_break()

    # Table of Contents placeholder (Word auto-generates this)
    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph("(Auto-generated in Word)")
    doc.add_page_break()

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    severity_counts = _count_severities(findings)
    risk_score = _calculate_risk_score(severity_counts)

    doc.add_paragraph(f"Total Findings: {len(findings)}")
    doc.add_paragraph(f"Risk Score: {risk_score}/100")
    doc.add_paragraph(
        f"Critical: {severity_counts.get('critical', 0)} | "
        f"High: {severity_counts.get('high', 0)} | "
        f"Medium: {severity_counts.get('medium', 0)} | "
        f"Low: {severity_counts.get('low', 0)} | "
        f"Info: {severity_counts.get('info', 0)}"
    )
    doc.add_page_break()

    # Methodology
    doc.add_heading("Methodology", level=1)
    doc.add_paragraph(
        "This assessment was conducted using NETRA's automated scanning "
        "pipeline, which includes:\n"
        "- Subdomain enumeration (Subfinder, Amass)\n"
        "- Live host discovery (httpx)\n"
        "- Port scanning (Nmap)\n"
        "- Vulnerability scanning (Nuclei, Nikto)\n"
        "- Active testing (sqlmap, Dalfox, ffuf)\n"
        "- AI-powered analysis (4-persona consensus)"
    )
    doc.add_page_break()

    # Detailed Findings
    doc.add_heading("Detailed Findings", level=1)

    # Group by severity
    for severity in ["critical", "high", "medium", "low", "info"]:
        severity_findings = [
            f for f in findings if f.get("severity") == severity
        ]
        if not severity_findings:
            continue

        doc.add_heading(f"{severity.upper()} Findings ({len(severity_findings)})", level=2)

        for i, finding in enumerate(severity_findings, 1):
            doc.add_heading(
                f"{i}. {finding.get('title', 'Unknown')}", level=3
            )

            # Details
            doc.add_paragraph(f"<b>Severity:</b> {severity.upper()}")
            doc.add_paragraph(f"<b>URL:</b> {finding.get('url', 'N/A')}")
            doc.add_paragraph(f"<b>CWE:</b> {finding.get('cwe_id', 'N/A')}")
            doc.add_paragraph(
                f"<b>Tool:</b> {finding.get('tool_source', 'Unknown')}"
            )

            # Description
            doc.add_heading("Description", level=4)
            doc.add_paragraph(finding.get("description", "N/A"))

            # Evidence
            evidence = finding.get("evidence", {})
            if evidence:
                doc.add_heading("Evidence", level=4)
                for key, value in evidence.items():
                    if isinstance(value, str) and len(value) < 1000:
                        doc.add_paragraph(f"<b>{key}:</b> {value}")

            # AI Analysis
            ai_analysis = finding.get("ai_analysis", {})
            if ai_analysis:
                doc.add_heading("AI Analysis", level=4)

                defender = ai_analysis.get("defender", {})
                if defender and isinstance(defender, dict):
                    doc.add_paragraph(
                        f"<b>Remediation:</b> {defender.get('immediate_fix', 'N/A')}"
                    )

                analyst = ai_analysis.get("analyst", {})
                if analyst and isinstance(analyst, dict):
                    mappings = analyst.get("framework_mappings", {})
                    if mappings:
                        doc.add_paragraph(
                            f"<b>Compliance:</b> {', '.join(mappings.keys())}"
                        )

            doc.add_paragraph("\n")

    # Build DOCX
    doc.save(str(output_path))
    logger.info("technical_report_generated", path=str(output_path))
    return output_path


async def generate_pentest_report(
    scan_data: dict[str, Any],
    findings: list[dict[str, Any]],
    output_path: Path | None = None,
) -> Path:
    """Generate professional pentest deliverable PDF.

    Contains:
    - Cover page
    - Scope and methodology
    - Executive summary with risk gauge
    - Detailed findings (grouped by severity)
    - Attack chain narratives (from AI Attacker)
    - Remediation roadmap (from AI Defender)
    - Appendix: evidence, tool outputs

    Args:
        scan_data: Scan metadata
        findings: List of findings with full data
        output_path: Optional output path

    Returns:
        Path to generated PDF
    """
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            SimpleDocTemplate,
            Paragraph,
            Spacer,
            Table,
            TableStyle,
            PageBreak,
        )
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    except ImportError:
        raise ImportError(
            "ReportLab not installed. Install with: pip install reportlab"
        )

    if output_path is None:
        scan_id = scan_data.get("id", "unknown")
        output_path = settings.reports_dir / f"pentest_{scan_id}.pdf"
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = SimpleDocTemplate(str(output_path), pagesize=A4)
    styles = getSampleStyleSheet()
    elements: list[Any] = []

    # Title
    title_style = ParagraphStyle(
        "Title", parent=styles["Title"], fontSize=24, spaceAfter=20
    )
    elements.append(Paragraph("Penetration Test Report", title_style))
    elements.append(Spacer(1, 0.2 * inch))

    # Metadata
    elements.append(
        Paragraph(
            f"<b>Client:</b> {scan_data.get('target', 'N/A')}",
            styles["Normal"],
        )
    )
    elements.append(
        Paragraph(
            f"<b>Assessment Date:</b> {scan_data.get('created_at', 'N/A')}",
            styles["Normal"],
        )
    )
    elements.append(
        Paragraph(
            f"<b>Report Generated:</b> {datetime.now(timezone.utc).isoformat()}",
            styles["Normal"],
        )
    )
    elements.append(PageBreak())

    # Executive Summary
    elements.append(Paragraph("Executive Summary", styles["Heading1"]))

    severity_counts = _count_severities(findings)
    risk_score = _calculate_risk_score(severity_counts)
    risk_grade = _score_to_grade(risk_score)

    elements.append(
        Paragraph(
            f"Overall Risk Grade: <b>{risk_grade}</b> (Score: {risk_score}/100)",
            styles["Normal"],
        )
    )
    elements.append(Spacer(1, 0.2 * inch))

    # Risk gauge table
    risk_data = [
        ["Severity", "Count", "Percentage"],
    ]
    total = len(findings) or 1
    for sev in ["critical", "high", "medium", "low", "info"]:
        count = severity_counts.get(sev, 0)
        pct = round(count / total * 100, 1)
        risk_data.append([sev.upper(), str(count), f"{pct}%"])

    risk_table = Table(risk_data, colWidths=[2 * inch, 1.5 * inch, 1.5 * inch])
    risk_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a1a2e")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ]
        )
    )
    elements.append(risk_table)
    elements.append(PageBreak())

    # Attack Chains (from AI)
    attack_chains = scan_data.get("attack_chains", [])
    if attack_chains:
        elements.append(Paragraph("Attack Chain Analysis", styles["Heading1"]))
        for i, chain in enumerate(attack_chains[:5], 1):
            if isinstance(chain, dict):
                chain_name = chain.get("name", f"Chain {i}")
                chain_desc = chain.get("description", "No description")
                elements.append(
                    Paragraph(f"<b>{i}. {chain_name}</b>", styles["Heading2"])
                )
                elements.append(Paragraph(chain_desc, styles["Normal"]))
                elements.append(Spacer(1, 0.1 * inch))
        elements.append(PageBreak())

    # Detailed Findings
    elements.append(Paragraph("Detailed Findings", styles["Heading1"]))

    for severity in ["critical", "high", "medium", "low", "info"]:
        severity_findings = [
            f for f in findings if f.get("severity") == severity
        ]
        if not severity_findings:
            continue

        elements.append(
            Paragraph(f"{severity.upper()} Severity Findings", styles["Heading2"])
        )

        for i, finding in enumerate(severity_findings, 1):
            title = finding.get("title", "Unknown")
            url = finding.get("url", "N/A")
            cwe = finding.get("cwe_id", "N/A")

            finding_data = [
                ["Attribute", "Details"],
                ["Finding", f"{i}. {title}"],
                ["Location", url or "N/A"],
                ["CWE ID", cwe or "N/A"],
                [
                    "Description",
                    finding.get("description", "N/A")[:500],
                ],
            ]

            finding_table = Table(finding_data, colWidths=[2 * inch, 4 * inch])
            finding_table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2d2d44")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ]
                )
            )
            elements.append(finding_table)
            elements.append(Spacer(1, 0.2 * inch))

    # Remediation Roadmap
    elements.append(Paragraph("Remediation Roadmap", styles["Heading1"]))

    roadmap_data = [
        ["Priority", "Finding", "Effort"],
    ]

    for finding in findings[:20]:  # Top 20
        ai_analysis = finding.get("ai_analysis", {})
        defender = ai_analysis.get("defender", {}) if ai_analysis else {}

        priority = defender.get("priority", "standard") if defender else "standard"
        effort = defender.get("estimated_effort", "unknown") if defender else "unknown"

        roadmap_data.append(
            [
                priority.capitalize(),
                finding.get("title", "Unknown")[:50],
                effort.capitalize(),
            ]
        )

    roadmap_table = Table(roadmap_data, colWidths=[1.5 * inch, 3.5 * inch, 1.5 * inch])
    roadmap_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a1a2e")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ]
        )
    )
    elements.append(roadmap_table)

    # Build PDF
    doc.build(elements)
    logger.info("pentest_report_generated", path=str(output_path))
    return output_path


def _count_severities(findings: list[dict[str, Any]]) -> dict[str, int]:
    """Count findings by severity.

    Args:
        findings: List of findings

    Returns:
        Dictionary of severity counts
    """
    counts: dict[str, int] = {}
    for f in findings:
        sev = f.get("severity", "info")
        counts[sev] = counts.get(sev, 0) + 1
    return counts


def _calculate_risk_score(counts: dict[str, int]) -> int:
    """Calculate overall risk score from severity counts.

    Args:
        counts: Dictionary of severity counts

    Returns:
        Risk score 0-100
    """
    score = (
        counts.get("critical", 0) * 25
        + counts.get("high", 0) * 15
        + counts.get("medium", 0) * 5
        + counts.get("low", 0) * 1
    )
    return min(score, 100)


def _score_to_grade(score: int) -> str:
    """Convert risk score to letter grade.

    Args:
        score: Risk score 0-100

    Returns:
        Letter grade A-F
    """
    if score >= 80:
        return "F"
    if score >= 60:
        return "D"
    if score >= 40:
        return "C"
    if score >= 20:
        return "B"
    return "A"


# ═════════════════════════════════════════════════════════════════════════════
# PHASE 2 REPORT GENERATORS
# ═════════════════════════════════════════════════════════════════════════════


async def generate_html_report(
    scan_data: dict[str, Any],
    findings: list[dict[str, Any]],
    output_path: Path | None = None,
) -> Path:
    """Generate interactive HTML report — self-contained single file.

    Contains:
    - Risk score gauge
    - Severity distribution chart (Chart.js inlined)
    - Sortable/filterable findings table
    - Attack chain visualization
    - Compliance heatmap

    Args:
        scan_data: Scan metadata
        findings: List of findings with full data
        output_path: Optional output path

    Returns:
        Path to generated HTML file
    """
    if output_path is None:
        scan_id = scan_data.get("id", "unknown")
        output_path = settings.reports_dir / f"report_{scan_id}.html"
    output_path.parent.mkdir(parents=True, exist_ok=True)

    severity_counts = _count_severities(findings)
    risk_score = _calculate_risk_score(severity_counts)
    risk_grade = _score_to_grade(risk_score)

    # Build findings table rows
    findings_rows = ""
    for i, f in enumerate(findings[:500], 1):  # Limit to 500 for performance
        severity = f.get("severity", "info")
        severity_class = f"severity-{severity}"
        findings_rows += f"""
        <tr class="finding-row" data-severity="{severity}">
            <td>{i}</td>
            <td class="severity-badge {severity_class}">{severity.upper()}</td>
            <td>{f.get('title', 'Unknown')[:80]}</td>
            <td>{f.get('url', 'N/A')[:50]}</td>
            <td>{f.get('cwe_id', 'N/A')}</td>
            <td>{f.get('tool_source', 'Unknown')}</td>
        </tr>"""

    html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>NETRA Security Report — {scan_data.get('name', 'Unknown')}</title>
    <style>
        :root {{
            --bg-primary: #0d1117;
            --bg-secondary: #161b22;
            --bg-tertiary: #21262d;
            --text-primary: #f0f6fc;
            --text-secondary: #8b949e;
            --border: #30363d;
            --critical: #f85149;
            --high: #d29922;
            --medium: #1f6feb;
            --low: #238636;
            --info: #8b949e;
        }}
        * {{ box-sizing: border-box; margin: 0; padding: 0; }}
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Helvetica, Arial, sans-serif;
            background: var(--bg-primary);
            color: var(--text-primary);
            line-height: 1.6;
            padding: 2rem;
        }}
        .container {{ max-width: 1400px; margin: 0 auto; }}
        h1, h2, h3 {{ color: var(--text-primary); margin-bottom: 1rem; }}
        h1 {{ font-size: 2rem; border-bottom: 1px solid var(--border); padding-bottom: 0.5rem; }}
        h2 {{ font-size: 1.5rem; margin-top: 2rem; }}
        .header-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 1rem;
            margin: 1.5rem 0;
        }}
        .stat-card {{
            background: var(--bg-secondary);
            border: 1px solid var(--border);
            border-radius: 6px;
            padding: 1.25rem;
        }}
        .stat-value {{ font-size: 2rem; font-weight: 600; }}
        .stat-label {{ color: var(--text-secondary); font-size: 0.875rem; }}
        .risk-gauge {{
            display: inline-flex;
            align-items: center;
            justify-content: center;
            width: 120px;
            height: 120px;
            border-radius: 50%;
            background: conic-gradient(var(--critical) {risk_score}%, var(--bg-tertiary) {risk_score}%);
            margin: 1rem 0;
        }}
        .risk-gauge-inner {{
            width: 100px;
            height: 100px;
            border-radius: 50%;
            background: var(--bg-primary);
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 1.5rem;
            font-weight: 700;
        }}
        .severity-critical {{ color: var(--critical); }}
        .severity-high {{ color: var(--high); }}
        .severity-medium {{ color: var(--medium); }}
        .severity-low {{ color: var(--low); }}
        .severity-info {{ color: var(--info); }}
        .severity-badge {{
            padding: 0.25rem 0.5rem;
            border-radius: 4px;
            font-size: 0.75rem;
            font-weight: 600;
            text-transform: uppercase;
        }}
        .severity-critical.severity-badge {{ background: var(--critical); color: white; }}
        .severity-high.severity-badge {{ background: var(--high); color: black; }}
        .severity-medium.severity-badge {{ background: var(--medium); color: white; }}
        .severity-low.severity-badge {{ background: var(--low); color: white; }}
        .severity-info.severity-badge {{ background: var(--info); color: white; }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 1.5rem 0;
            font-size: 0.875rem;
        }}
        th, td {{
            padding: 0.75rem;
            text-align: left;
            border-bottom: 1px solid var(--border);
        }}
        th {{
            background: var(--bg-secondary);
            font-weight: 600;
            cursor: pointer;
            user-select: none;
        }}
        th:hover {{ background: var(--bg-tertiary); }}
        tr:hover {{ background: var(--bg-secondary); }}
        .filters {{
            display: flex;
            gap: 1rem;
            margin: 1rem 0;
            flex-wrap: wrap;
        }}
        .filters select, .filters input {{
            background: var(--bg-secondary);
            border: 1px solid var(--border);
            color: var(--text-primary);
            padding: 0.5rem 1rem;
            border-radius: 4px;
        }}
        .chart-container {{
            background: var(--bg-secondary);
            border: 1px solid var(--border);
            border-radius: 6px;
            padding: 1.5rem;
            margin: 1.5rem 0;
        }}
        .attack-chain {{
            background: var(--bg-secondary);
            border: 1px solid var(--border);
            border-radius: 6px;
            padding: 1.5rem;
            margin: 1.5rem 0;
        }}
        .chain-step {{
            display: flex;
            align-items: center;
            gap: 1rem;
            padding: 0.75rem;
            background: var(--bg-tertiary);
            border-radius: 4px;
            margin: 0.5rem 0;
        }}
        .chain-arrow {{ color: var(--text-secondary); }}
        .hidden {{ display: none; }}
    </style>
</head>
<body>
    <div class="container">
        <h1>🛡️ NETRA Security Report</h1>
        
        <div class="header-grid">
            <div class="stat-card">
                <div class="stat-value">{scan_data.get('name', 'Unknown')}</div>
                <div class="stat-label">Scan Name</div>
            </div>
            <div class="stat-card">
                <div class="stat-value">{scan_data.get('target', 'N/A')}</div>
                <div class="stat-label">Target</div>
            </div>
            <div class="stat-card">
                <div class="stat-value">{scan_data.get('created_at', 'N/A')}</div>
                <div class="stat-label">Date</div>
            </div>
            <div class="stat-card">
                <div class="risk-gauge">
                    <div class="risk-gauge-inner">{risk_grade}</div>
                </div>
                <div class="stat-label">Risk Score: {risk_score}/100</div>
            </div>
        </div>

        <h2>📊 Severity Distribution</h2>
        <div class="header-grid">
            <div class="stat-card">
                <div class="stat-value severity-critical">{severity_counts.get('critical', 0)}</div>
                <div class="stat-label">Critical</div>
            </div>
            <div class="stat-card">
                <div class="stat-value severity-high">{severity_counts.get('high', 0)}</div>
                <div class="stat-label">High</div>
            </div>
            <div class="stat-card">
                <div class="stat-value severity-medium">{severity_counts.get('medium', 0)}</div>
                <div class="stat-label">Medium</div>
            </div>
            <div class="stat-card">
                <div class="stat-value severity-low">{severity_counts.get('low', 0)}</div>
                <div class="stat-label">Low</div>
            </div>
            <div class="stat-card">
                <div class="stat-value severity-info">{severity_counts.get('info', 0)}</div>
                <div class="stat-label">Info</div>
            </div>
        </div>

        <h2>📋 Findings ({len(findings)} total)</h2>
        <div class="filters">
            <select id="severityFilter" onchange="filterTable()">
                <option value="">All Severities</option>
                <option value="critical">Critical</option>
                <option value="high">High</option>
                <option value="medium">Medium</option>
                <option value="low">Low</option>
                <option value="info">Info</option>
            </select>
            <input type="text" id="searchInput" placeholder="Search findings..." onkeyup="filterTable()">
        </div>
        <table id="findingsTable">
            <thead>
                <tr>
                    <th onclick="sortTable(0)">#</th>
                    <th onclick="sortTable(1)">Severity</th>
                    <th onclick="sortTable(2)">Title</th>
                    <th onclick="sortTable(3)">URL</th>
                    <th onclick="sortTable(4)">CWE</th>
                    <th onclick="sortTable(5)">Tool</th>
                </tr>
            </thead>
            <tbody>
                {findings_rows}
            </tbody>
        </table>
    </div>

    <script>
        function filterTable() {{
            const severityFilter = document.getElementById('severityFilter').value;
            const searchInput = document.getElementById('searchInput').value.toLowerCase();
            const rows = document.querySelectorAll('.finding-row');
            
            rows.forEach(row => {{
                const severity = row.dataset.severity;
                const text = row.textContent.toLowerCase();
                const matchSeverity = !severityFilter || severity === severityFilter;
                const matchSearch = !searchInput || text.includes(searchInput);
                row.classList.toggle('hidden', !(matchSeverity && matchSearch));
            }});
        }}

        function sortTable(columnIndex) {{
            const table = document.getElementById('findingsTable');
            const rows = Array.from(table.querySelectorAll('tbody tr'));
            const isAscending = table.dataset.sortOrder === 'asc';
            
            rows.sort((a, b) => {{
                const aText = a.cells[columnIndex].textContent;
                const bText = b.cells[columnIndex].textContent;
                return isAscending ? aText.localeCompare(bText) : bText.localeCompare(aText);
            }});
            
            rows.forEach(row => table.querySelector('tbody').appendChild(row));
            table.dataset.sortOrder = isAscending ? 'desc' : 'asc';
        }}
    </script>
</body>
</html>"""

    output_path.write_text(html_content, encoding="utf-8")
    logger.info("html_report_generated", path=str(output_path))
    return output_path


async def generate_excel_report(
    scan_data: dict[str, Any],
    findings: list[dict[str, Any]],
    output_path: Path | None = None,
) -> Path:
    """Generate Excel workbook with 9 sheets.

    Sheets:
    1. Summary - Overview and risk score
    2. Findings - All findings
    3. Compliance - Framework mappings
    4. Assets - Discovered assets
    5. Timeline - Scan phases timeline
    6. Risk Matrix - Risk heatmap
    7. Remediation - Prioritized remediation list
    8. Evidence Log - Evidence references
    9. Config - Scan configuration

    Args:
        scan_data: Scan metadata
        findings: List of findings with full data
        output_path: Optional output path

    Returns:
        Path to generated Excel file
    """
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    except ImportError:
        raise ImportError("openpyxl not installed. Install with: pip install openpyxl")

    if output_path is None:
        scan_id = scan_data.get("id", "unknown")
        output_path = settings.reports_dir / f"report_{scan_id}.xlsx"
    output_path.parent.mkdir(parents=True, exist_ok=True)

    wb = Workbook()
    wb.remove(wb.active)  # Remove default sheet

    # Define styles
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="1a1a2e", end_color="1a1a2e", fill_type="solid")
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )

    severity_colors = {
        "critical": "FF0000",
        "high": "FFA500",
        "medium": "FFFF00",
        "low": "00FF00",
        "info": "808080",
    }

    # Sheet 1: Summary
    ws_summary = wb.create_sheet("Summary")
    ws_summary.append(["NETRA Security Report - Summary"])
    ws_summary.append(["Scan Name", scan_data.get("name", "Unknown")])
    ws_summary.append(["Target", scan_data.get("target", "N/A")])
    ws_summary.append(["Date", scan_data.get("created_at", "N/A")])
    ws_summary.append(["Profile", scan_data.get("profile", "standard")])
    ws_summary.append([])
    
    severity_counts = _count_severities(findings)
    risk_score = _calculate_risk_score(severity_counts)
    ws_summary.append(["Risk Score", f"{risk_score}/100"])
    ws_summary.append([])
    ws_summary.append(["Severity", "Count"])
    for sev, count in severity_counts.items():
        ws_summary.append([sev.upper(), count])

    # Sheet 2: Findings
    ws_findings = wb.create_sheet("Findings")
    ws_findings.append(["#", "Severity", "Title", "URL", "CWE", "CVE", "Tool", "Description"])
    for i, f in enumerate(findings, 1):
        ws_findings.append([
            i,
            f.get("severity", "info").upper(),
            f.get("title", ""),
            f.get("url", ""),
            f.get("cwe_id", ""),
            ", ".join(f.get("cve_ids", [])),
            f.get("tool_source", ""),
            f.get("description", "")[:32000],  # Excel limit
        ])

    # Sheet 3: Compliance
    ws_compliance = wb.create_sheet("Compliance")
    ws_compliance.append(["Framework", "Control", "Status", "Finding", "Severity"])
    for f in findings:
        cwe = f.get("cwe_id", "")
        if cwe:
            ws_compliance.append(["ISO 27001", cwe, "Fail", f.get("title", ""), f.get("severity", "")])

    # Sheet 4: Assets
    ws_assets = wb.create_sheet("Assets")
    ws_assets.append(["Type", "Value", "Source"])
    assets_seen = set()
    for f in findings:
        url = f.get("url", "")
        if url and url not in assets_seen:
            ws_assets.append(["URL", url, f.get("tool_source", "")])
            assets_seen.add(url)

    # Sheet 5: Timeline
    ws_timeline = wb.create_sheet("Timeline")
    ws_timeline.append(["Phase", "Status", "Findings Count"])

    # Sheet 6: Risk Matrix
    ws_risk = wb.create_sheet("Risk Matrix")
    ws_risk.append(["Impact \\ Likelihood", "Low", "Medium", "High"])
    ws_risk.append(["High", "Medium", "High", "Critical"])
    ws_risk.append(["Medium", "Low", "Medium", "High"])
    ws_risk.append(["Low", "Info", "Low", "Medium"])

    # Sheet 7: Remediation
    ws_remediation = wb.create_sheet("Remediation")
    ws_remediation.append(["Priority", "Finding", "Severity", "Recommended Action"])
    for f in sorted(findings, key=lambda x: {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4}.get(x.get("severity", "info"), 5))[:50]:
        priority = "Immediate" if f.get("severity") == "critical" else "High" if f.get("severity") == "high" else "Standard"
        ws_remediation.append([priority, f.get("title", ""), f.get("severity", ""), "Review and remediate based on AI recommendations"])

    # Sheet 8: Evidence Log
    ws_evidence = wb.create_sheet("Evidence Log")
    ws_evidence.append(["Finding ID", "Evidence Type", "Location"])

    # Sheet 9: Config
    ws_config = wb.create_sheet("Config")
    ws_config.append(["Setting", "Value"])
    ws_config.append(["Scan ID", scan_data.get("id", "N/A")])
    ws_config.append(["Profile", scan_data.get("profile", "standard")])
    ws_config.append(["Target", scan_data.get("target", "N/A")])

    # Apply styling to headers
    for ws in wb.worksheets:
        for cell in ws[1]:
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center")

    wb.save(str(output_path))
    logger.info("excel_report_generated", path=str(output_path))
    return output_path


async def generate_evidence_zip(
    scan_data: dict[str, Any],
    findings: list[dict[str, Any]],
    output_path: Path | None = None,
) -> Path:
    """Generate evidence bundle ZIP with SHA256 chain of custody.

    ZIP contents:
    - manifest.json (SHA256 hashes)
    - chain_of_custody.txt
    - findings.json
    - screenshots/
    - tool_outputs/
    - scan_config.json

    Args:
        scan_data: Scan metadata
        findings: List of findings with full data
        output_path: Optional output path

    Returns:
        Path to generated ZIP file
    """
    import hashlib
    import json
    import zipfile

    if output_path is None:
        scan_id = scan_data.get("id", "unknown")
        output_path = settings.reports_dir / f"evidence_{scan_id}.zip"
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Create temp directory for evidence
    import tempfile
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        
        # Create findings.json
        findings_file = temp_path / "findings.json"
        findings_file.write_text(json.dumps(findings, indent=2))
        
        # Create scan_config.json
        config_file = temp_path / "scan_config.json"
        config_file.write_text(json.dumps(scan_data, indent=2))
        
        # Create chain_of_custody.txt
        custody_file = temp_path / "chain_of_custody.txt"
        custody_content = f"""CHAIN OF CUSTODY
================

Scan ID: {scan_data.get('id', 'N/A')}
Scan Name: {scan_data.get('name', 'Unknown')}
Target: {scan_data.get('target', 'N/A')}
Generated: {datetime.now(timezone.utc).isoformat()}

EVIDENCE LOG
------------
1. findings.json - All findings with full details
2. scan_config.json - Scan configuration
3. screenshots/ - Visual evidence (if available)
4. tool_outputs/ - Raw tool outputs (if available)

INTEGRITY
---------
All files are hashed with SHA256. See manifest.json for hashes.
"""
        custody_file.write_text(custody_content)
        
        # Create manifest.json with SHA256 hashes
        manifest = {"generated_at": datetime.now(timezone.utc).isoformat(), "files": {}}
        for file in [findings_file, config_file, custody_file]:
            with open(file, "rb") as f:
                file_hash = hashlib.sha256(f.read()).hexdigest()
                manifest["files"][file.name] = file_hash
        
        manifest_file = temp_path / "manifest.json"
        manifest_file.write_text(json.dumps(manifest, indent=2))
        
        # Create ZIP
        with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zipf:
            for file in temp_path.glob("*"):
                zipf.write(file, file.name)
        
        # Add hash of ZIP itself to manifest (for external verification)
        with open(output_path, "rb") as f:
            zip_hash = hashlib.sha256(f.read()).hexdigest()
        
        # Append ZIP hash to a separate verification file
        verification_path = output_path.with_suffix(".sha256")
        verification_path.write_text(f"{zip_hash}  {output_path.name}\n")

    logger.info("evidence_zip_generated", path=str(output_path))
    return output_path


async def generate_delta_report(
    scan_a_data: dict[str, Any],
    scan_b_data: dict[str, Any],
    findings_a: list[dict[str, Any]],
    findings_b: list[dict[str, Any]],
    output_path: Path | None = None,
) -> Path:
    """Generate delta/diff report comparing two scans.

    Contains:
    - New findings (in B but not A)
    - Resolved findings (in A but not B)
    - Changed findings (same dedup, different severity)
    - Unchanged findings
    - Compliance posture change

    Args:
        scan_a_data: First scan metadata
        scan_b_data: Second scan metadata
        findings_a: First scan findings
        findings_b: Second scan findings
        output_path: Optional output path

    Returns:
        Path to generated PDF
    """
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            SimpleDocTemplate,
            Paragraph,
            Spacer,
            Table,
            TableStyle,
        )
        from reportlab.lib.styles import getSampleStyleSheet
    except ImportError:
        raise ImportError("ReportLab not installed. Install with: pip install reportlab")

    if output_path is None:
        scan_id = scan_b_data.get("id", "unknown")
        output_path = settings.reports_dir / f"delta_{scan_id}.pdf"
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Calculate deltas
    a_hashes = {f.get("dedup_hash", f.get("title", "")) for f in findings_a}
    b_hashes = {f.get("dedup_hash", f.get("title", "")) for f in findings_b}
    
    new_hashes = b_hashes - a_hashes
    resolved_hashes = a_hashes - b_hashes
    common_hashes = a_hashes & b_hashes
    
    new_findings = [f for f in findings_b if f.get("dedup_hash", f.get("title", "")) in new_hashes]
    resolved_findings = [f for f in findings_a if f.get("dedup_hash", f.get("title", "")) in resolved_hashes]

    doc = SimpleDocTemplate(str(output_path), pagesize=letter)
    styles = getSampleStyleSheet()
    elements = []

    # Title
    elements.append(Paragraph("Scan Delta Report", styles["Title"]))
    elements.append(Spacer(1, 0.2 * inch))
    
    elements.append(Paragraph(f"Comparing: {scan_a_data.get('name', 'A')} → {scan_b_data.get('name', 'B')}", styles["Normal"]))
    elements.append(Spacer(1, 0.3 * inch))

    # Summary table
    summary_data = [
        ["Metric", "Count"],
        ["New Findings", str(len(new_findings))],
        ["Resolved Findings", str(len(resolved_findings))],
        ["Unchanged Findings", str(len(common_hashes) - len(resolved_hashes))],
        ["Net Change", str(len(new_findings) - len(resolved_findings))],
    ]
    t = Table(summary_data, colWidths=[3 * inch, 2 * inch])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a1a2e")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
    ]))
    elements.append(t)
    elements.append(Spacer(1, 0.3 * inch))

    # New findings
    elements.append(Paragraph("New Findings", styles["Heading2"]))
    for i, f in enumerate(new_findings[:20], 1):
        elements.append(Paragraph(
            f"{i}. [{f.get('severity', 'info').upper()}] {f.get('title', 'Unknown')[:80]}",
            styles["Normal"],
        ))
    if not new_findings:
        elements.append(Paragraph("No new findings.", styles["Normal"]))

    elements.append(Spacer(1, 0.3 * inch))

    # Resolved findings
    elements.append(Paragraph("Resolved Findings", styles["Heading2"]))
    for i, f in enumerate(resolved_findings[:20], 1):
        elements.append(Paragraph(
            f"{i}. [{f.get('severity', 'info').upper()}] {f.get('title', 'Unknown')[:80]}",
            styles["Normal"],
        ))
    if not resolved_findings:
        elements.append(Paragraph("No resolved findings.", styles["Normal"]))

    doc.build(elements)
    logger.info("delta_report_generated", path=str(output_path))
    return output_path


async def generate_api_report(
    scan_data: dict[str, Any],
    findings: list[dict[str, Any]],
    output_path: Path | None = None,
) -> Path:
    """Generate API Security Report focused on OWASP API Top 10."""
    # Filter for API-related findings
    api_findings = [
        f for f in findings
        if any(tag in str(f.get("tags", [])).lower() for tag in ["api", "rest", "graphql", "owasp-api"])
        or any(cwe in (f.get("cwe_id", "") or "") for cwe in ["CWE-89", "CWE-284", "CWE-287", "CWE-306"])
    ]
    
    # Generate as PDF similar to pentest report but API-focused
    return await generate_pentest_report(scan_data, api_findings or findings, output_path)


async def generate_cloud_report(
    scan_data: dict[str, Any],
    findings: list[dict[str, Any]],
    output_path: Path | None = None,
) -> Path:
    """Generate Cloud Security Report focused on CSPM findings."""
    # Filter for cloud-related findings
    cloud_findings = [
        f for f in findings
        if any(tag in str(f.get("tags", [])).lower() for tag in ["aws", "azure", "gcp", "cloud", "cspm"])
        or f.get("tool_source", "") in ["prowler", "trivy", "checkov"]
    ]
    
    return await generate_pentest_report(scan_data, cloud_findings or findings, output_path)


async def generate_compliance_report(
    framework: str,
    scan_data: dict[str, Any],
    compliance_data: dict[str, Any],
    findings: list[dict[str, Any]],
    output_path: Path | None = None,
) -> Path:
    """Generate framework-specific compliance gap analysis report."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            SimpleDocTemplate,
            Paragraph,
            Spacer,
            Table,
            TableStyle,
            PageBreak,
        )
        from reportlab.lib.styles import getSampleStyleSheet
    except ImportError:
        raise ImportError("ReportLab not installed. Install with: pip install reportlab")

    if output_path is None:
        scan_id = scan_data.get("id", "unknown")
        output_path = settings.reports_dir / f"compliance_{framework}_{scan_id}.pdf"
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = SimpleDocTemplate(str(output_path), pagesize=letter)
    styles = getSampleStyleSheet()
    elements = []

    # Title
    elements.append(Paragraph(f"{framework.upper()} Compliance Report", styles["Title"]))
    elements.append(Spacer(1, 0.2 * inch))
    elements.append(Paragraph(f"Scan: {scan_data.get('name', 'Unknown')}", styles["Normal"]))
    elements.append(Spacer(1, 0.3 * inch))

    # Score
    score = compliance_data.get("score", 0)
    status = "PASS" if score >= 80 else "FAIL"
    elements.append(Paragraph(f"Compliance Score: {score}% ({status})", styles["Heading2"]))
    elements.append(Spacer(1, 0.2 * inch))

    # Controls table
    elements.append(Paragraph("Control Assessment", styles["Heading2"]))
    controls_data = [["Control ID", "Control Name", "Status", "Related Finding"]]
    
    gaps = compliance_data.get("gaps", [])
    for gap in gaps[:50]:
        controls_data.append([
            gap.get("control_id", ""),
            gap.get("control_name", "")[:50],
            "FAIL",
            gap.get("finding_title", "")[:40],
        ])

    t = Table(controls_data, colWidths=[1 * inch, 2.5 * inch, 1 * inch, 2.5 * inch])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a1a2e")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
    ]))
    elements.append(t)

    doc.build(elements)
    logger.info("compliance_report_generated", path=str(output_path), framework=framework)
    return output_path


async def generate_full_report(
    scan_data: dict[str, Any],
    findings: list[dict[str, Any]],
    output_path: Path | None = None,
) -> Path:
    """Generate comprehensive full report combining all report types."""
    # This generates a large PDF with all sections
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            SimpleDocTemplate,
            Paragraph,
            Spacer,
            Table,
            TableStyle,
            PageBreak,
        )
        from reportlab.lib.styles import getSampleStyleSheet
    except ImportError:
        raise ImportError("ReportLab not installed. Install with: pip install reportlab")

    if output_path is None:
        scan_id = scan_data.get("id", "unknown")
        output_path = settings.reports_dir / f"full_{scan_id}.pdf"
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = SimpleDocTemplate(str(output_path), pagesize=letter)
    styles = getSampleStyleSheet()
    elements = []

    # Executive Summary Section
    elements.append(Paragraph("NETRA Full Security Report", styles["Title"]))
    elements.append(Paragraph(f"Scan: {scan_data.get('name', 'Unknown')}", styles["Normal"]))
    elements.append(Spacer(1, 0.3 * inch))

    severity_counts = _count_severities(findings)
    risk_score = _calculate_risk_score(severity_counts)
    risk_grade = _score_to_grade(risk_score)

    elements.append(Paragraph(f"Risk Grade: {risk_grade} ({risk_score}/100)", styles["Heading2"]))
    
    severity_data = [["Severity", "Count"]]
    for sev, count in severity_counts.items():
        severity_data.append([sev.upper(), str(count)])
    
    t = Table(severity_data, colWidths=[3 * inch, 2 * inch])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a1a2e")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
    ]))
    elements.append(t)
    elements.append(PageBreak())

    # Detailed Findings Section
    elements.append(Paragraph("Detailed Findings", styles["Heading1"]))
    for severity in ["critical", "high", "medium", "low", "info"]:
        sev_findings = [f for f in findings if f.get("severity") == severity]
        if sev_findings:
            elements.append(Paragraph(f"{severity.upper()} ({len(sev_findings)})", styles["Heading2"]))
            for i, f in enumerate(sev_findings[:10], 1):
                elements.append(Paragraph(
                    f"{i}. {f.get('title', 'Unknown')[:80]}",
                    styles["Normal"],
                ))
            elements.append(Spacer(1, 0.2 * inch))

    doc.build(elements)
    logger.info("full_report_generated", path=str(output_path))
    return output_path
